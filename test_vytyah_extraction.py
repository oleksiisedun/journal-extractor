"""
Mini test of a local LLM for generating ЖБД (combat log) extracts.

See README.md for setup/run instructions and CLAUDE.md for full pipeline
and rule documentation.
"""

import json
import os
import re
from datetime import date

import requests
from docx import Document

# ---------- SETTINGS ----------
OLLAMA_URL = "http://localhost:11434/api/chat"
MODEL = "qwen3:8b-q8_0"          # switch to "qwen3:8b" for faster/lighter
ZHBD_PATH = "journals/ЖБД_02_04_2026.docx"  # path to the daily ЖБД file

# ---------- STEP 1: Parse the source (no LLM) ----------

def load_paragraphs(docx_path):
    """Extracts all non-empty paragraphs with continuous numbering, verbatim
    (no pandoc/reflow)."""
    doc = Document(docx_path)
    result = []
    idx = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        result.append((idx, p.text))
                        idx += 1
    return result


def load_paragraph_columns(docx_path):
    """Column-aware companion to load_paragraphs(), for the time-of-day
    extraction heuristic (see assign_time_boundaries()). Walks the same
    table/row/cell structure in the same order, so global indices for
    column-1 (content) paragraphs match load_paragraphs()'s numbering
    exactly and can be cross-referenced against an LLM pointer's
    target_paragraph_index. Also keeps, per row, each paragraph's RAW
    position within its own cell (including blank paragraphs) for both
    columns -- the time column visually lines up with wrapped content text
    via blank filler paragraphs, and that raw position is what the
    proportional cut-point math in assign_time_boundaries() needs.

    Assumes column 0 = time, column 1 = content, which holds for the one
    real sample file available so far (journals/ЖБД_02_04_2026.docx) --
    re-validate against other ЖБД templates if the layout ever differs.
    """
    doc = Document(docx_path)
    rows_out = []
    idx = 0
    for table in doc.tables:
        for row in table.rows:
            time_labels = []
            content_paragraphs = []
            time_raw_count = len(row.cells[0].paragraphs) if row.cells else 0
            content_raw_count = len(row.cells[1].paragraphs) if len(row.cells) > 1 else 0
            for col, cell in enumerate(row.cells):
                for raw_idx, p in enumerate(cell.paragraphs):
                    if not p.text.strip():
                        continue
                    if col == 0:
                        time_labels.append((raw_idx, p.text))
                    elif col == 1:
                        content_paragraphs.append((idx, raw_idx, p.text))
                    idx += 1
            rows_out.append({
                "time_labels": time_labels,              # [(raw_col0_idx, text)]
                "time_raw_count": time_raw_count,         # len(cell[0].paragraphs), incl. blanks
                "content_paragraphs": content_paragraphs, # [(global_idx, raw_col1_idx, text)]
                "content_raw_count": content_raw_count,   # len(cell[1].paragraphs), incl. blanks
            })
    return rows_out


FILENAME_DATE_PATTERN = re.compile(r"(\d{2})[_.\-](\d{2})[_.\-](\d{4})")


def extract_date_from_filename(docx_path):
    """Extracts the day's date from a ЖБД filename (e.g. 'ЖБД_02_04_2026.docx',
    'ЖБД 10.07.2026.docx', 'ЖБД_12-04-2026.docx' -> 2026-04-02 / 2026-07-10 /
    2026-04-12, Ukrainian day-first convention). The separator between the
    DD/MM/YYYY groups varies by file (seen so far: '_', '.', '-') and isn't
    required to be consistent within one filename. Fails loud (ValueError)
    rather than silently returning no date, since a wrong or missing date
    would misattribute every fragment extracted from this file."""
    basename = os.path.basename(docx_path)
    match = FILENAME_DATE_PATTERN.search(basename)
    if not match:
        raise ValueError(
            f"Could not find a DD_MM_YYYY date in filename: {basename!r} "
            f"(expected a pattern like 'ЖБД_02_04_2026.docx')"
        )
    day, month, year = (int(g) for g in match.groups())
    try:
        return date(year, month, day)
    except ValueError as e:
        raise ValueError(
            f"Filename {basename!r} encodes an invalid calendar date "
            f"(day={day}, month={month}, year={year}): {e}"
        ) from e


def find_candidate_windows(paragraphs, surname, window=8):
    """Deterministic prefilter (no LLM): paragraph windows around mentions
    of the SURNAME, not the full name — avoids losing namesakes at this
    stage."""
    para_dict = dict(paragraphs)
    max_idx = max(para_dict.keys())
    hits = [i for i, t in paragraphs if surname.upper() in t.upper()]
    if not hits:
        return []

    raw_windows = [(max(0, i - window), min(max_idx, i + 1)) for i in hits]
    raw_windows.sort()
    merged = [raw_windows[0]]
    for lo, hi in raw_windows[1:]:
        if lo <= merged[-1][1] + 1:
            merged[-1] = (merged[-1][0], max(merged[-1][1], hi))
        else:
            merged.append((lo, hi))
    return merged


def extract_full_name(rank_and_name):
    """Surname + first name + patronymic (without rank) — for narrower
    candidate filtering than searching by surname alone."""
    tokens = rank_and_name.split()
    surname_idx = None
    for i, w in enumerate(tokens):
        letters = [c for c in w if c.isalpha()]
        if letters and w == w.upper():
            surname_idx = i
            break
    if surname_idx is None:
        raise ValueError(f"Could not determine surname in: {rank_and_name!r}")
    return " ".join(tokens[surname_idx:surname_idx + 3])


def filter_windows_by_full_name(paragraphs, windows, full_name):
    """Narrows surname-based windows down to the ones where the FULL name
    occurs verbatim, removing ambiguity before the LLM call rather than
    relying on the model to pick the right window."""
    para_dict = dict(paragraphs)
    matched = []
    for lo, hi in windows:
        text = " ".join(para_dict[i] for i in range(lo, hi + 1))
        if full_name.upper() in text.upper():
            matched.append((lo, hi))
    return matched


# ---------- STEP 2: LLM query (only "where", never "what") ----------

SYSTEM_PROMPT = """You locate the mention of ONE specific person in a numbered list of
paragraphs from a single day of a ЖБД (military combat log).

You are given: rank + full name of the person to find, and a numbered list
of paragraphs.

Output THREE things, never the text itself:
1. "context_paragraph_indices" — paragraphs needed for legal/section context:
   order references (containing "БР", "наказ", "розпорядження", or a number
   like №.../.../ДСК), section headers (e.g. "на ПВ «...»:"). These do NOT
   have to be adjacent to the target paragraph — there can be a large gap
   with OTHER people's own paragraphs in between, and that gap is simply
   skipped, not included. Context paragraphs must NEVER contain other
   people's names.
2. "target_paragraph_index" — the SINGLE paragraph that contains the target
   person. Match the FULL name (surname + first name + patronymic) — the
   document may contain several people with the same surname.
3. "redactions" — ONLY used when OTHER people appear INSIDE that SAME target
   paragraph (e.g. several names in one continuous sentence). Copy their
   text VERBATIM, character-for-character. This is RARE — most of the time
   redactions stays empty, because other people are on their OWN separate
   paragraphs and are simply never selected (see Example 2 below).

   NEVER put the target person's own name into redactions.
   NEVER try to redact a paragraph that is not context_paragraph_indices or
   target_paragraph_index — if another person is on their own separate
   paragraph, just don't select that paragraph. Do not enumerate or redact
   people who are not on the target's own paragraph.

   CRITICAL: the paragraph list you receive may contain candidates from
   MULTIPLE, UNRELATED sections/orders (e.g. two different people with the
   same surname, each governed by a DIFFERENT order). Pick context
   paragraphs ONLY from the SAME section/order that actually governs the
   target paragraph. NEVER combine order-reference paragraphs from two
   different, unrelated orders into context_paragraph_indices — a single
   person's record is governed by exactly ONE order, never two.

If the person is not found in the given list, return found: false,
context_paragraph_indices: [], target_paragraph_index: -1, redactions: [].

CRITICAL: you NEVER rewrite or paraphrase any text. Your job is only to
point to WHERE things are, not to write text yourself.

--- EXAMPLE 1: target is the first person right after a section header ---
[36] на виконання БР командира 7 обр (на бмп) 2 боп 15 обр «СОКІЛ» №БР42/Б3/7Р/ДСК від 15.02.26, продовжують стійко утримувати зайняті рубежі оборони в районі ВЕРБОВЕ, СОСНІВКА, Полтавської громади, Полтавського району, Полтавської області, не допускають раптових дії противника на лінії бойового зіткнення з противником в першому ешелоні оборони:
[37] на ПВ «БЕРЕГ»:
[38] солдат ОРЛЕНКО Олександр Сергійович;
[39] солдат ПЕТРЕНКО Едуард Дмитрович;

Query: "солдат ОРЛЕНКО Олександр Сергійович"
Correct answer:
{"found": true, "context_paragraph_indices": [36, 37], "target_paragraph_index": 38, "redactions": []}
(Paragraph 39 — a different person — is simply never selected. No redaction needed.)

--- EXAMPLE 2: target is deep inside a long list, with OTHER people's own
paragraphs in between the context and the target ---
[89] на виконання БР командира 7 обр (на бмп) 2 боп 15 обр «СОКІЛ» №БР47/Б3/7Р/ДСК від 15.02.26, з метою швидкого реагування ...
[90] молодший сержант КОВАЛЬЧУК Євген Миколайович;
[91] сержант ЮЩЕНКО Ренат Маратович;
[92] старший солдат ДЕМЧЕНКО Олександр Васильович;
[93] солдат ОРЛЕНКО Павло Юрійович;
[94] солдат МАЛЬЦЕВА Олександра Васильович;

Query: "солдат ОРЛЕНКО Павло Юрійович"
Correct answer:
{"found": true, "context_paragraph_indices": [89], "target_paragraph_index": 93, "redactions": []}
(Paragraphs 90-92 and 94 — other people, each on their own paragraph — are
simply never selected and need NO redactions. Do NOT try to list them in
redactions — that is unnecessary and error-prone. Only the context
paragraph [89] and the target paragraph [93] are chosen.)

--- EXAMPLE 3: TWO people share the SAME paragraph — this is when redactions
is actually needed ---
[55] з метою вивчення реального стану справ ... проводили роботу: головний сержант 2 корпусу НГУ «СОКІЛ» майстер-сержант БОНДАРЕНКО Дмитро Сергійович, водій ... солдат ЛИТВИНЕНКО Андрій Іванович (згідно БР ... від 10.02.2026).

Query: "солдат ЛИТВИНЕНКО Андрій Іванович"
Correct answer:
{"found": true, "context_paragraph_indices": [], "target_paragraph_index": 55, "redactions": ["головний сержант 2 корпусу НГУ «СОКІЛ» майстер-сержант БОНДАРЕНКО Дмитро Сергійович, "]}
(Here БОНДАРЕНКО shares the exact same paragraph/sentence as the target, so his
text must be redacted. This is the ONLY situation where redactions is used.)

--- EXAMPLE 4 (rare fallback case): if the list still contains candidates
from TWO DIFFERENT, UNRELATED orders (e.g. two namesakes governed by
different orders) — pick context from the order that actually governs the
target paragraph, never both. A person is governed by exactly ONE order."""

RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "found": {"type": "boolean"},
        "context_paragraph_indices": {
            "type": "array",
            "items": {"type": "integer"},
            "description": "Paragraphs needed for legal/section context (order references, "
                           "section headers). NOT necessarily adjacent to the target paragraph, "
                           "and must NEVER contain other people's names.",
        },
        "target_paragraph_index": {"type": "integer"},
        "redactions": {
            "type": "array",
            "items": {"type": "string"},
            "description": "Verbatim text of OTHER people to remove, ONLY if they appear "
                           "INSIDE the target paragraph itself (same sentence).",
        },
    },
    "required": ["found", "context_paragraph_indices", "target_paragraph_index", "redactions"],
}


def ask_llm(paragraphs, rank_and_name):
    """Sends the candidate paragraph window to the local LLM and returns the
    parsed pointer JSON — never the assembled text itself."""
    listing = "\n".join(f"[{i}] {text}" for i, text in paragraphs)
    user_prompt = (
        f"Looking for: {rank_and_name}\n\n"
        f"Paragraph list (may be a candidate window, not the full day):\n{listing}"
    )
    payload = {
        "model": MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        "format": RESPONSE_SCHEMA,
        "stream": False,
        "think": False,          # CRITICAL: top-level, not inside "options" —
                                  # otherwise Qwen3 (a hybrid thinking model)
                                  # still generates a full reasoning chain, and
                                  # on CPU that is several times slower. The
                                  # "/no_think" text trick in the prompt is an
                                  # unreliable legacy method — don't rely on it.
        "keep_alive": "30m",     # keep the model loaded in memory between calls
        "options": {
            "temperature": 0,
            "num_thread": 8,      # physical cores of the Ryzen 7 7840HS (not
                                   # 16 — SMT often doesn't speed up, and can
                                   # slow down, memory-bound inference; verify
                                   # 16 yourself too)
        },
    }
    resp = requests.post(OLLAMA_URL, json=payload, timeout=300)
    resp.raise_for_status()
    content = resp.json()["message"]["content"]
    return json.loads(content)


# ---------- STEPS 3-5: Deterministic assembly + validation (no LLM) ----------

# Optional space after № tolerated — real files are inconsistent about it
# (e.g. '№БР42/Б3/7Р/ДСК' vs '№ БР91/Б3/12Р/ДСК').
ORDER_REF_PATTERN = re.compile(r"№\s?\S+")

# MGRS-style grid coordinate, e.g. "37U CR 1234 5678" — zone+band, 100km
# square, then two equal-length easting/northing digit groups.
COORDINATE_PATTERN = re.compile(r"\d{1,2}[A-Z]\s[A-Z]{2}\s\d{2,5}\s\d{2,5}")

# "район зосередження" (concentration area) and its grammatical variants
# (районі, району, зосередженню, ...), with an immediately preceding
# preposition (у/в/на) swallowed too so no dangling preposition is left.
LOCATION_LABEL_PATTERN = re.compile(
    r"(?:[ув]|на)\s+район\w*\s+зосередженн\w*|район\w*\s+зосередженн\w*",
    re.IGNORECASE,
)


def extract_order_refs(text):
    """Extracts order/directive numbers of the form №БР42/Б3/7Р/ДСК from
    text. Whitespace after № is stripped before comparison so the same
    order isn't treated as two different ones just because one mention
    happens to have a space after № and another doesn't (both forms occur
    in real files)."""
    return {re.sub(r"\s+", "", ref) for ref in ORDER_REF_PATTERN.findall(text)}


# ---------- Time-of-day extraction (no LLM) ----------
#
# Two formats seen in real files, apparently never mixed within one day:
# (a) "easy" / inline — a content paragraph itself starts with the time
#     (e.g. '00.00 на виконання БОЙОВОГО РОЗПОРЯДЖЕННЯ ...'), and the time
#     column is empty. This is exact/verbatim-derived — no guessing needed.
# (b) "hard" / left-column — the narrow left column carries the times, NOT
#     paragraph-aligned with the content column at all — verified
#     empirically: gaps between consecutive left-column time labels don't
#     track the content column's structure proportionally. A naive
#     proportional mapping lands correctly on real section/list headers
#     some of the time, but for the rest it lands INSIDE a single,
#     homogeneous list of people governed by one order — which would
#     silently mis-assign a time to real people. Handled with a
#     snap-to-nearest-boundary heuristic below; only used when (a) finds
#     nothing in a row. Assumes column 0 = time, column 1 = content.

ROMAN_HEADER_PATTERN = re.compile(r"^[IVXІ]+\.")

INLINE_TIME_PATTERN = re.compile(r"^(\d{2}\.\d{2}(?:-\d{2}\.\d{2})?)\b")


def extract_inline_time(text):
    """Returns the time token at the very start of a content paragraph
    (e.g. '00.00' from '00.00 на виконання ...', or '05.00-06.40' from a
    ranged form), or None if the paragraph doesn't start with one. This is
    the "easy case": exact, verbatim-derived, no heuristic guessing."""
    match = INLINE_TIME_PATTERN.match(text.strip())
    return match.group(1) if match else None


def is_boundary_paragraph(text):
    """A content paragraph is a plausible time-boundary point if it's a
    Roman-numeral section header, a list/section-intro line ending in ':'
    (e.g. 'на ПВ «БРАВО»:'), or an order-reference sentence (matches
    ORDER_REF_PATTERN). Ordinary list entries (a single person's line)
    never qualify — that's what stops a time label from being snapped into
    the middle of a person list."""
    stripped = text.strip()
    return bool(
        ROMAN_HEADER_PATTERN.match(stripped)
        or stripped.endswith(":")
        or ORDER_REF_PATTERN.search(stripped)
    )


def _assign_time_boundaries_inline(content_paragraphs):
    """The "easy case": content paragraphs that themselves start with a
    time token. Exact and verbatim-derived, so always "confident"."""
    return [
        (global_idx, extract_inline_time(text), "confident")
        for global_idx, raw_idx, text in sorted(content_paragraphs, key=lambda t: t[0])
        if extract_inline_time(text)
    ]


def _assign_time_boundaries_left_column(row, slack):
    """The "hard case": snap-to-nearest-boundary heuristic over the left
    time column. See the module-level comment above for why this can't be
    an exact lookup. For each time label, a proportional cut point is
    computed from its raw position in the (blank-padded) time column, then
    snapped FORWARD to the nearest boundary-classified content paragraph.
    A label is "uncertain" if that snap traveled more than `slack` raw
    paragraphs from its predicted cut point, or if it collided with the
    same boundary as the previous label (in which case only the later,
    more specific time is kept — we can't tell which of the two genuinely
    governs that point)."""
    time_labels = row["time_labels"]
    left_raw_count = row["time_raw_count"]
    content_paragraphs = row["content_paragraphs"]
    right_raw_count = row["content_raw_count"]
    if not time_labels or left_raw_count <= 1 or right_raw_count <= 1:
        return []

    boundary_candidates = sorted(
        (raw_idx, global_idx)
        for global_idx, raw_idx, text in content_paragraphs
        if is_boundary_paragraph(text)
    )

    raw_assignments = []
    for label_raw_idx, time_value in time_labels:
        cut_point = round(
            label_raw_idx / (left_raw_count - 1) * (right_raw_count - 1)
        )
        snapped = next((b for b in boundary_candidates if b[0] >= cut_point), None)
        if snapped is None:
            continue  # no boundary left to snap to — this label can't be placed
        snapped_raw_idx, snapped_global_idx = snapped
        confident = abs(snapped_raw_idx - cut_point) <= slack
        raw_assignments.append([snapped_global_idx, time_value, confident])

    collapsed = []
    for global_idx, time_value, confident in raw_assignments:
        if collapsed and collapsed[-1][0] == global_idx:
            # two labels snapped to the same boundary — ambiguous which one
            # actually governs it; keep the later (more specific) time
            collapsed[-1] = [global_idx, time_value, False]
        else:
            collapsed.append([global_idx, time_value, confident])

    return [
        (global_idx, time_value, "confident" if confident else "uncertain")
        for global_idx, time_value, confident in collapsed
    ]


def assign_time_boundaries(row, slack=15):
    """Deterministically maps a table row's time information onto its
    content-column paragraphs. Returns an ascending list of
    (global_content_index, time_value, confidence) tuples, where
    confidence is "confident" or "uncertain" — never silently asserts a
    time it isn't reasonably sure about (same posture as the surname/
    order-number guardrails in assemble_fragment()).

    Tries the inline ("easy case") format first — see
    _assign_time_boundaries_inline() — and only falls back to the
    left-column heuristic (_assign_time_boundaries_left_column()) if no
    inline time tokens were found in this row's content. The two formats
    are assumed mutually exclusive within a single row/day, matching every
    real file seen so far; re-check this assumption if a file ever mixes
    them.
    """
    inline = _assign_time_boundaries_inline(row["content_paragraphs"])
    if inline:
        return inline
    return _assign_time_boundaries_left_column(row, slack)


def time_for_paragraph(boundaries, target_global_idx):
    """Looks up the (time_value, confidence) governing a given global
    content paragraph index: the boundary assignment with the largest
    global_content_index <= target_global_idx. Returns None if the target
    occurs before any boundary was assigned (time genuinely unknown)."""
    result = None
    for global_idx, time_value, confidence in boundaries:
        if global_idx <= target_global_idx:
            result = (time_value, confidence)
        else:
            break
    return result


def strip_coordinates(text):
    """Removes MGRS-style grid coordinates (e.g. '37U CR 1234 5678'), including
    semicolon-separated lists of them, from the given text. Coordinates are
    tactical/operational data and must never appear in the generated
    extract, regardless of how the source formats them. Also drops any
    parenthetical group that becomes empty once its coordinates are removed
    (e.g. 'район зосередження (37U CR 1234 5678; 37U CR 1234 5678)' ->
    'район зосередження')."""
    text = COORDINATE_PATTERN.sub("", text)

    def _drop_if_empty(match):
        inner = match.group(1)
        return "" if re.fullmatch(r"[\s;]*", inner) else match.group(0)

    text = re.sub(r"\(([^()]*)\)", _drop_if_empty, text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"[ \t]+([,;.):])", r"\1", text)
    text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)
    return text


def strip_location_labels(text):
    """Removes the phrase 'район зосередження' and its grammatical variants
    (районі, району, зосередженню, ...) from text, swallowing an immediately
    preceding preposition (у/в/на) so no dangling preposition is left. Like
    coordinates, this is location/tactical information with no place in a
    personnel extract — applied unconditionally, not just when it looks
    out of place."""
    text = LOCATION_LABEL_PATTERN.sub("", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"[ \t]+([,;.):])", r"\1", text)
    text = re.sub(r"^[ \t]+", "", text, flags=re.MULTILINE)
    text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)
    return text


def assemble_fragment(paragraphs, pointer, date_value=None, time_boundaries=None):
    """Deterministically assembles the final fragment from the LLM's pointer
    — slices source text, applies redactions, runs guardrails, applies the
    one allowed punctuation fix. Also attaches the day's date (from the
    filename, via extract_date_from_filename()) and the heuristically
    assigned time (via assign_time_boundaries() + time_for_paragraph()) —
    both deterministic, no LLM involvement. Returns a dict, not a bare
    string, since date/time metadata rides along with the text."""
    para_dict = dict(paragraphs)

    if not pointer["found"]:
        return None

    context_indices = pointer["context_paragraph_indices"]
    target_index = pointer["target_paragraph_index"]
    indices = sorted(set(context_indices + [target_index]))
    parts = [para_dict[i] for i in indices]

    # SANITY GUARDRAIL: if TWO DIFFERENT order numbers appear among the
    # CONTEXT paragraphs (not among all selected paragraphs — the target
    # paragraph itself may not contain a number at all), that's a sign the
    # model stuck in context from an unrelated, foreign order (found
    # empirically: context=[36, 89], where 36 belongs to БР18 for a
    # DIFFERENT person, and 89 to the correct БР19). One person cannot be
    # governed by two different orders at once within a single record.
    order_refs = set()
    for i in context_indices:
        order_refs |= extract_order_refs(para_dict[i])
    if len(order_refs) > 1:
        raise ValueError(
            f"MODEL LOGIC ERROR: context paragraphs reference MULTIPLE "
            f"different orders at once ({order_refs}) — this is a sign the "
            f"model mixed in context from an unrelated, foreign section. "
            f"Result REJECTED.\ncontext_paragraph_indices={context_indices}"
        )

    # apply redactions — exact str.replace, fail-closed
    for r in pointer["redactions"]:
        applied = False
        for i in range(len(parts)):
            if r in parts[i]:
                parts[i] = parts[i].replace(r, "")
                applied = True
        if not applied:
            raise ValueError(
                f"VALIDATION ERROR: string to remove was not found verbatim "
                f"in the selected paragraphs — result REJECTED, manual review needed:\n{r!r}"
            )

    # mechanical strip — coordinates and location labels carry no legal/
    # identity meaning for the extract and must never leak into the output
    # (see strip_coordinates, strip_location_labels)
    parts = [strip_location_labels(strip_coordinates(p)) for p in parts]

    fragment_text = "\n\n".join(parts)

    # SANITY GUARDRAIL: the target person's full name must physically be
    # present in the final fragment. If the model mistakenly put the
    # target THEMSELVES into redactions (confusing "who is the target" with
    # "who to remove"), the fragment would come out without them, and this
    # must be caught here, not let through.
    surname = pointer.get("_surname_check")
    if surname and surname.upper() not in fragment_text.upper():
        raise ValueError(
            f"MODEL LOGIC ERROR: surname '{surname}' is missing from the "
            f"assembled fragment — the LLM likely redacted the target person "
            f"themselves by mistake. Result REJECTED.\nfragment_text={fragment_text!r}"
        )

    # the one allowed exception to the "verbatim" rule
    stripped = fragment_text.rstrip()
    if stripped.endswith(";"):
        stripped = stripped[:-1] + "."

    time_result = time_for_paragraph(time_boundaries, target_index) if time_boundaries else None
    time_value, time_confidence = time_result if time_result else (None, "uncertain")

    return {
        "text": stripped,
        "date": date_value,
        "time": time_value,
        "time_confidence": time_confidence,
    }


# ---------- TEST RUN ----------

def extract_surname(rank_and_name):
    """The surname is normally written in UPPERCASE — take the first such word."""
    for word in rank_and_name.split():
        letters = [c for c in word if c.isalpha()]
        if letters and word == word.upper():
            return word
    raise ValueError(f"Could not determine surname in: {rank_and_name!r}")


if __name__ == "__main__":
    all_paragraphs = load_paragraphs(ZHBD_PATH)
    print(f"Завантажено параграфів: {len(all_paragraphs)}\n")

    day_date = extract_date_from_filename(ZHBD_PATH)

    # find the table row that actually holds the day's content (the one
    # with the most content paragraphs) and derive its time boundaries —
    # see assign_time_boundaries() for why this is heuristic/best-effort
    columns = load_paragraph_columns(ZHBD_PATH)
    content_row = max(columns, key=lambda r: len(r["content_paragraphs"]))
    time_boundaries = assign_time_boundaries(content_row)
    print(f"Дата: {day_date.isoformat()}    Знайдено часових міток: {len(time_boundaries)}\n")

    test_cases = [
        "солдат ОРЛЕНКО Олександр Сергійович",   # present in the document (ПВ «БЕРЕГ»)
        "солдат ОРЛЕНКО Павло Юрійович",          # namesake with a different first name (reserve) — disambiguation test
        "сержант НЕІСНУЮЧИЙ Іван Іванович",        # found=false test
    ]

    for person in test_cases:
        print("=" * 70)
        print("Шукаємо:", person)

        surname = extract_surname(person)
        windows = find_candidate_windows(all_paragraphs, surname)

        if not windows:
            # the surname isn't in the day's text at all — no need to call
            # the LLM, and this is faster and more reliable than asking the
            # model "does this person exist"
            print(">>> Прізвище відсутнє в тексті дня — LLM не викликаємо. found=false")
            continue

        # narrow by FULL name (not surname alone) - if this unambiguously
        # narrows to a single window, the LLM never sees other namesakes
        # or other orders at all. If not (0 or 2+ windows), fall back to
        # giving the LLM everything found by surname, as before
        # (guardrails already exist for this case).
        full_name = extract_full_name(person)
        narrowed = filter_windows_by_full_name(all_paragraphs, windows, full_name)
        if len(narrowed) == 1:
            windows_to_use = narrowed
            note = "однозначно звужено до 1 вікна за повним ПІБ"
        elif len(narrowed) == 0:
            windows_to_use = windows
            note = "ПОВНЕ ПІБ не знайдено дослівно (можлива описка?) - даємо LLM усі вікна за прізвищем"
        else:
            windows_to_use = narrowed
            note = f"увага: повне ПІБ дослівно зустрічається у {len(narrowed)} різних місцях - справжня неоднозначність"

        candidate_paragraphs = []
        seen = set()
        for lo, hi in windows_to_use:
            for i in range(lo, hi + 1):
                if i not in seen:
                    seen.add(i)
                    candidate_paragraphs.append((i, dict(all_paragraphs)[i]))
        print(f"    (префільтр: {len(candidate_paragraphs)} з {len(all_paragraphs)} параграфів, "
              f"{note})")

        pointer = ask_llm(candidate_paragraphs, person)
        pointer["_surname_check"] = surname  # for the sanity check in assemble_fragment
        print("Відповідь LLM (pointer):", pointer)

        if not pointer.get("found"):
            print(">>> Не знайдено в цьому дні (перевір вручну — гепи не пропускаємо мовчки)")
            continue

        result = assemble_fragment(
            all_paragraphs, pointer,
            date_value=day_date, time_boundaries=time_boundaries,
        )
        time_note = "" if result["time_confidence"] == "confident" else "  [!! ЧАС НЕВИЗНАЧЕНИЙ/НЕТОЧНИЙ — перевірити вручну !!]"
        print(f">>> Дата: {result['date'].isoformat()}    Час: {result['time']}{time_note}")
        print(">>> ЗІБРАНИЙ ФРАГМЕНТ (дослівно з джерела):")
        print(result["text"])
        print()
