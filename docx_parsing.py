"""Parses a day's .docx combat log into indexed, verbatim paragraph lists
(no LLM involvement)."""

import os
import re
from datetime import date

from docx import Document


def load_paragraphs(docx_path):
    """Extracts all non-empty paragraphs with continuous numbering, verbatim
    (no pandoc/reflow)."""
    doc = Document(docx_path)
    result = []
    idx = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        result.append((idx, p.text))
                        idx += 1
    return result


def load_paragraph_columns(docx_path):
    """Column-aware companion to load_paragraphs(), for the time-of-day
    extraction heuristic (see assign_time_boundaries()). Walks the same
    table/row/cell structure in the same order, so global indices for
    column-1 (content) paragraphs match load_paragraphs()'s numbering
    exactly and can be cross-referenced against an LLM pointer's
    target_paragraph_index. Also keeps, per row, each paragraph's RAW
    position within its own cell (including blank paragraphs) for both
    columns -- the time column visually lines up with wrapped content text
    via blank filler paragraphs, and that raw position is what the
    proportional cut-point math in assign_time_boundaries() needs.

    Assumes column 0 = time, column 1 = content, which holds for the one
    real sample file available so far (journals/ЖБД_02_04_2026.docx) --
    re-validate against other combat log templates if the layout ever differs.
    """
    doc = Document(docx_path)
    rows_out = []
    idx = 0
    for table in doc.tables:
        for row in table.rows:
            time_labels = []
            content_paragraphs = []
            time_raw_count = len(row.cells[0].paragraphs) if row.cells else 0
            content_raw_count = len(row.cells[1].paragraphs) if len(row.cells) > 1 else 0
            for col, cell in enumerate(row.cells):
                for raw_idx, p in enumerate(cell.paragraphs):
                    if not p.text.strip():
                        continue
                    if col == 0:
                        time_labels.append((raw_idx, p.text))
                    elif col == 1:
                        content_paragraphs.append((idx, raw_idx, p.text))
                    idx += 1
            rows_out.append({
                "time_labels": time_labels,              # [(raw_col0_idx, text)]
                "time_raw_count": time_raw_count,         # len(cell[0].paragraphs), incl. blanks
                "content_paragraphs": content_paragraphs, # [(global_idx, raw_col1_idx, text)]
                "content_raw_count": content_raw_count,   # len(cell[1].paragraphs), incl. blanks
            })
    return rows_out


FILENAME_DATE_PATTERN = re.compile(r"(\d{2})[_.\-](\d{2})[_.\-](\d{4})")


def extract_date_from_filename(docx_path):
    """Extracts the day's date from a combat log filename (e.g. 'ЖБД_02_04_2026.docx',
    'ЖБД 10.07.2026.docx', 'ЖБД_12-04-2026.docx' -> 2026-04-02 / 2026-07-10 /
    2026-04-12, Ukrainian day-first convention). The separator between the
    DD/MM/YYYY groups varies by file (seen so far: '_', '.', '-') and isn't
    required to be consistent within one filename. Fails loud (ValueError)
    rather than silently returning no date, since a wrong or missing date
    would misattribute every fragment extracted from this file."""
    basename = os.path.basename(docx_path)
    match = FILENAME_DATE_PATTERN.search(basename)
    if not match:
        raise ValueError(
            f"Could not find a DD_MM_YYYY date in filename: {basename!r} "
            f"(expected a pattern like 'ЖБД_02_04_2026.docx')"
        )
    day, month, year = (int(g) for g in match.groups())
    try:
        return date(year, month, day)
    except ValueError as e:
        raise ValueError(
            f"Filename {basename!r} encodes an invalid calendar date "
            f"(day={day}, month={month}, year={year}): {e}"
        ) from e
