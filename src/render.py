"""Fills templates/1.docx-style extract templates with assembled fragments
(no LLM, no text generation — see CLAUDE.md's "Core architectural
principle"). Every placeholder is replaced with text that already came
verbatim out of assembly.assemble_fragment(); this module only arranges
that text into the template's paragraphs, it never invents any of it.

Callers are expected to have already run merge.merge_consecutive_entries()
on the chronological "found" days, so `entries` here carries "date_from"/
"date_to" (equal for a single unmerged day) rather than a bare "date".
"""

import copy
import os

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from text_wrap import estimate_wrapped_line_count

TWIPS_PER_POINT = 20

# OOXML's built-in default table cell margin (start/end), used only when
# neither the cell's own tcMar nor the table's tblCellMar specify one.
DEFAULT_CELL_MARGIN_TWIPS = 108

# Word's built-in default run font size, used only if templates/1.docx's
# {витяг} run, paragraph mark, and style chain all somehow specify none.
DEFAULT_FONT_SIZE_PT = 10.0


def _replace_run_placeholder(document, placeholder, value):
    """Substring-replaces `placeholder` inside whichever run of a top-level
    (non-table) paragraph contains it, e.g. '3 ДСК/Х/7 від {дата витягу}'
    -> '...від 01.08.2026' — formatting of the surrounding run is
    untouched since only its text content changes."""
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                return
    raise ValueError(f"Placeholder {placeholder!r} not found in template")


def _find_placeholder_cell_and_paragraph(document, placeholder):
    """Finds the table cell and paragraph whose entire text is exactly
    `placeholder` (e.g. '{дата}', '{витяг}') — in the real template each
    one is the sole content of its own paragraph, distinct from
    _replace_run_placeholder's case of a placeholder sharing a run with
    other text. Returns the cell too (not just the paragraph) so callers
    can inspect/adjust the paragraphs preceding it, e.g.
    _equalize_leading_blanks()."""
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text == placeholder:
                        return cell, paragraph
    raise ValueError(f"Placeholder {placeholder!r} not found in template")


def _placeholder_prefix_length(cell, placeholder_paragraph):
    """Counts how many paragraphs precede `placeholder_paragraph` at the
    start of `cell` — its static, pre-existing template content (e.g. the
    {витяг} cell's fixed "V. Хід бойових дій" header + blank line).
    Compares the stable underlying `_p` XML element rather than the
    `Paragraph` wrapper, since `cell.paragraphs` builds a fresh wrapper
    list on every access."""
    target = placeholder_paragraph._p
    count = 0
    for paragraph in cell.paragraphs:
        if paragraph._p is target:
            break
        count += 1
    return count


def _remove_leading_blanks(cell, count):
    """Deletes up to `count` blank paragraphs from the start of `cell`,
    stopping early if a non-blank paragraph is reached first — so real
    static template content (e.g. the {витяг} cell's "V. Хід бойових дій"
    header) is never deleted, even if `count` asks for more than the
    cell's actual blank-only prefix."""
    removed = 0
    for paragraph in list(cell.paragraphs):
        if removed >= count or paragraph.text != "":
            break
        p_element = paragraph._p
        p_element.getparent().remove(p_element)
        removed += 1


def _equalize_leading_blanks(date_cell, date_paragraph, fragment_cell, fragment_paragraph):
    """The real template has a different number of static leading
    paragraphs before {дата} (3 blanks) than before {витяг} (1 "V. Хід
    бойових дій" header line + 1 blank) — 3 vs 2. Even after
    _format_date_lines() pads each entry's date block to match its text
    block's line count, this constant prefix mismatch alone leaves the two
    columns' per-entry content starting at different paragraph offsets, so
    dates still render a fixed number of lines below their matching text.
    Trims whichever cell has the longer prefix down to the other's length
    (only ever deleting blank paragraphs — see _remove_leading_blanks — so
    real static text is never touched), so _expand_multiline_placeholder()
    inserts both cells' entries starting at the same offset."""
    date_prefix = _placeholder_prefix_length(date_cell, date_paragraph)
    fragment_prefix = _placeholder_prefix_length(fragment_cell, fragment_paragraph)
    if date_prefix > fragment_prefix:
        _remove_leading_blanks(date_cell, date_prefix - fragment_prefix)
    elif fragment_prefix > date_prefix:
        _remove_leading_blanks(fragment_cell, fragment_prefix - date_prefix)


def _zero_space_after(p_element):
    """Sets a cloned paragraph's space-after to 0 twips — used for {дата}
    filler lines that stand in for one wrapped continuation line inside a
    single {витяг} paragraph. Word only charges space-after once per
    paragraph (after its last visual line), never once per wrapped line,
    so a run of filler paragraphs that each kept the placeholder's normal
    space-after would make the {дата} column taller than the {витяг} text
    it's supposed to line up with — see docs/bug-log.md."""
    pPr = p_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_element.insert(0, pPr)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:after"), "0")


def _expand_multiline_placeholder(paragraph, lines):
    """Replaces a paragraph whose only content is a placeholder run with
    one paragraph per (text, suppress_space_after) pair in `lines` (empty
    text renders as a blank paragraph, matching how real extract samples
    separate stacked entries). Each new paragraph is a deep copy of the
    placeholder paragraph's XML, so it inherits the exact same pPr/rPr
    formatting (font, size, spacing) without this code needing to know or
    hardcode any of it — except where `suppress_space_after` asks
    _zero_space_after() to override the cloned space-after to 0."""
    p_element = paragraph._p
    parent = p_element.getparent()

    for text, suppress_space_after in lines:
        new_p = copy.deepcopy(p_element)
        runs = new_p.findall(qn("w:r"))
        if text == "":
            for run_element in runs:
                new_p.remove(run_element)
        else:
            t_elements = new_p.findall(".//" + qn("w:t"))
            t_elements[0].text = text
            t_elements[0].set(qn("xml:space"), "preserve")
            for run_element in runs[1:]:
                new_p.remove(run_element)
        if suppress_space_after:
            _zero_space_after(new_p)
        p_element.addprevious(new_p)

    parent.remove(p_element)


def _format_time_line(entry):
    """One line for a single-day entry's time: the raw value as-is,
    confident or uncertain alike. Returns None when no time could be
    resolved at all, so the caller omits the line rather than presenting a
    fabricated value as extract text."""
    return entry["time"]


def _entry_fragment_lines(entry):
    """The {витяг} cell's paragraph lines for one entry — its already-
    assembled verbatim text split on the paragraph breaks assembly.py
    joined with "\\n"."""
    return entry["text"].split("\n")


def _cell_margin_twips(cell, side):
    """`cell`'s horizontal margin (side is "start" or "end") in twips: the
    cell's own tcMar override if present, else the enclosing table's
    default tblCellMar, else OOXML's built-in default
    (DEFAULT_CELL_MARGIN_TWIPS)."""
    tc = cell._tc
    tcPr = tc.tcPr
    if tcPr is not None:
        tcMar = tcPr.find(qn("w:tcMar"))
        if tcMar is not None:
            side_element = tcMar.find(qn(f"w:{side}"))
            if side_element is not None:
                return int(side_element.get(qn("w:w")))
    tbl = tc.getparent().getparent()  # w:tc -> w:tr -> w:tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblCellMar = tblPr.find(qn("w:tblCellMar"))
        if tblCellMar is not None:
            side_element = tblCellMar.find(qn(f"w:{side}"))
            if side_element is not None:
                return int(side_element.get(qn("w:w")))
    return DEFAULT_CELL_MARGIN_TWIPS


def _paragraph_first_line_indent_twips(paragraph):
    """`paragraph`'s first-line indent in twips (0 if unset) — the real
    template applies one to every {витяг} paragraph (a "red line" indent
    convention), which shrinks the usable width of a paragraph's first
    wrapped line relative to the lines after it."""
    pPr = paragraph._p.pPr
    if pPr is not None:
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            first_line = ind.get(qn("w:firstLine"))
            if first_line is not None:
                return int(first_line)
    return 0


def _run_font_size_pt(paragraph):
    """Font size in points used by `paragraph`'s text — its first run's
    explicit size if set (the case for the {витяг} placeholder, whose
    formatting _expand_multiline_placeholder() clones onto every inserted
    paragraph), else the paragraph mark's own run properties, else
    DEFAULT_FONT_SIZE_PT."""
    for run in paragraph.runs:
        if run.font.size is not None:
            return run.font.size.pt
    pPr = paragraph._p.pPr
    if pPr is not None:
        rPr = pPr.find(qn("w:rPr"))
        if rPr is not None:
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                return int(sz.get(qn("w:val"))) / 2
    return DEFAULT_FONT_SIZE_PT


def _fragment_line_widths_pt(fragment_cell, fragment_paragraph):
    """Returns (first_line_width_pt, continuation_width_pt) — the {витяг}
    cell's real usable text width for a paragraph's first line (reduced by
    its first-line indent) and for every line after, computed from
    templates/1.docx's actual cell width, margins, and indent rather than
    assumed."""
    usable_twips = (
        fragment_cell.width.twips
        - _cell_margin_twips(fragment_cell, "start")
        - _cell_margin_twips(fragment_cell, "end")
    )
    first_line_indent_twips = _paragraph_first_line_indent_twips(fragment_paragraph)
    first_line_width_pt = (usable_twips - first_line_indent_twips) / TWIPS_PER_POINT
    continuation_width_pt = usable_twips / TWIPS_PER_POINT
    return first_line_width_pt, continuation_width_pt


def _entry_visual_line_count(entry, font_size_pt, first_line_width_pt, continuation_width_pt):
    """Total visual lines the {витяг} cell will render for one entry — the
    sum of text_wrap.estimate_wrapped_line_count() over each of its
    paragraphs (_entry_fragment_lines), measured against the cell's real
    font size and usable width."""
    return sum(
        estimate_wrapped_line_count(line, font_size_pt, first_line_width_pt, continuation_width_pt)
        for line in _entry_fragment_lines(entry)
    )


def _format_date_lines(entries, font_size_pt, first_line_width_pt, continuation_width_pt):
    """Builds the {дата} cell's (text, suppress_space_after) pairs: for a
    single-day entry (date_from == date_to), the date plus its time line;
    for a merged multi-day range, one "з ... по ..." line and no time line
    — merge.merge_consecutive_entries() already dropped the time for
    ranges since no single time correctly represents the whole span. Each
    entry's date block is then padded with filler lines up to that same
    entry's *measured visual line count* in the {витяг} cell
    (_entry_visual_line_count) — not just its raw paragraph count, since a
    paragraph that wraps to several visual lines in Word still counts as
    one paragraph — so the two cells' entries start at matching visual
    offsets and the date lines up with the top of its own text block
    instead of landing early, inside a wrapped paragraph from an earlier
    entry.

    Filler lines are one docx paragraph each, so — unlike a wrapped
    continuation line inside a single {витяг} paragraph — every one of
    them would normally still charge the placeholder's paragraph
    space-after (see _zero_space_after's docstring and
    docs/bug-log.md). Left unchecked this makes the {дата} column
    taller than the {витяг} text it's supposed to track, and the drift
    compounds entry over entry. To match, only as many filler lines as
    the entry's real {витяг} *paragraph* count (_entry_fragment_lines)
    minus its own real content lines are allowed to keep normal
    space-after; the rest are marked for suppression. Which specific
    filler lines keep it doesn't affect the block's total height (a
    fixed amount of space-after is charged once per kept line,
    regardless of position) — only the count does — so the kept ones are
    simply placed first for simplicity.

    The last entry is never padded — that padding only exists to push a
    *later* entry's date down, and with no entry after it padding would
    just add trailing blank paragraphs that make the {дата} cell (and so
    the whole row) taller than the content needs, leaving a gap at the
    bottom of the table. Blank-line separated between entries (none
    trailing); the separator keeps normal space-after."""
    lines = []
    for i, entry in enumerate(entries):
        if i > 0:
            lines.append(("", False))
        entry_lines = []
        date_from = entry["date_from"].strftime("%d.%m.%Y")
        if entry["date_from"] == entry["date_to"]:
            entry_lines.append(date_from)
            time_line = _format_time_line(entry)
            if time_line is not None:
                entry_lines.append(time_line)
        else:
            date_to = entry["date_to"].strftime("%d.%m.%Y")
            entry_lines.append(f"з {date_from} по {date_to}")
        entry_pairs = [(line, False) for line in entry_lines]
        if i < len(entries) - 1:
            visual_line_count = _entry_visual_line_count(
                entry, font_size_pt, first_line_width_pt, continuation_width_pt
            )
            fragment_paragraph_count = len(_entry_fragment_lines(entry))
            filler_needed = max(0, visual_line_count - len(entry_lines))
            spaced_filler_needed = min(
                filler_needed, max(0, fragment_paragraph_count - len(entry_lines))
            )
            for j in range(filler_needed):
                entry_pairs.append(("", j >= spaced_filler_needed))
        lines.extend(entry_pairs)
    return lines


def _format_fragment_lines(entries):
    """Builds the {витяг} cell's (text, suppress_space_after) pairs: each
    entry's already-assembled verbatim text split on its own paragraph
    breaks (always normal space-after — suppression is only ever needed
    for {дата}'s filler lines), blank-line separated between entries (none
    trailing)."""
    lines = []
    for i, entry in enumerate(entries):
        if i > 0:
            lines.append(("", False))
        lines.extend((line, False) for line in _entry_fragment_lines(entry))
    return lines


def render_extract(entries, issue_date, template_path, output_path):
    """Fills `template_path` with `entries` (a chronological list of
    merge.merge_consecutive_entries()'s result dicts — {"text", "date_from",
    "date_to", "time", "time_confidence"} — for one person's "found" days
    only; callers are expected to have already filtered out "not_found"/
    "rejected" days and run the merge step) and saves the result to
    `output_path`. `issue_date` fills the header's {дата витягу}
    placeholder.
    """
    if not entries:
        raise ValueError("No fragments to render — entries is empty.")

    document = docx.Document(template_path)

    _replace_run_placeholder(document, "{дата витягу}", issue_date.strftime("%d.%m.%Y"))

    date_cell, date_paragraph = _find_placeholder_cell_and_paragraph(document, "{дата}")
    fragment_cell, fragment_paragraph = _find_placeholder_cell_and_paragraph(document, "{витяг}")
    _equalize_leading_blanks(date_cell, date_paragraph, fragment_cell, fragment_paragraph)

    font_size_pt = _run_font_size_pt(fragment_paragraph)
    first_line_width_pt, continuation_width_pt = _fragment_line_widths_pt(fragment_cell, fragment_paragraph)
    date_lines = _format_date_lines(entries, font_size_pt, first_line_width_pt, continuation_width_pt)

    _expand_multiline_placeholder(date_paragraph, date_lines)
    _expand_multiline_placeholder(fragment_paragraph, _format_fragment_lines(entries))

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    document.save(output_path)
